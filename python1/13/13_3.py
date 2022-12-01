import docx
import os

os.chdir('python1\\automate_online-materials')
doc = docx.Document('demo.docx')
print(len(doc.paragraphs))

def getText(fileName):
    doc = docx.Document(fileName)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)